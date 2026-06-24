import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert


OUTPUT_DIR = "output/tailored_resumes"


def _sanitize(text: str) -> str:
    return "".join(c if c.isalnum() or c in "-_" else "_" for c in text)


def build_resume_docx(resume: dict, job: dict) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    company = _sanitize(job.get("company", "company"))[:30]
    role = _sanitize(job.get("title", "role"))[:30]
    today = date.today().strftime("%Y%m%d")

    docx_path = os.path.join(OUTPUT_DIR, f"{company}_{role}_{today}.docx")
    pdf_path = os.path.join(OUTPUT_DIR, f"{company}_{role}_{today}.pdf")

    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    personal = resume.get("personal_info", {})
    name = personal.get("name", "")

    # Header — name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    # Contact line
    contact_parts = [
        personal.get("email", ""),
        personal.get("phone", ""),
        personal.get("linkedin", ""),
        personal.get("github", ""),
        personal.get("portfolio", ""),
    ]
    contact_line = " | ".join(p for p in contact_parts if p)
    if contact_line:
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)

    # Summary
    summary = resume.get("summary", "")
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    # Experience
    experience = resume.get("experience", [])
    if experience:
        doc.add_heading("Experience", level=1)
        for exp in experience:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('title', '')}  —  {exp.get('company', '')}")
            run.bold = True
            dates_para = doc.add_paragraph(exp.get("dates", ""))
            if dates_para.runs:
                dates_para.runs[0].italic = True
                dates_para.runs[0].font.size = Pt(10)
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # Skills
    skills = resume.get("skills", [])
    if skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(skills))

    # Education
    education = resume.get("education", [])
    if education:
        doc.add_heading("Education", level=1)
        for edu in education:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')}  —  {edu.get('institution', '')}")
            run.bold = True
            doc.add_paragraph(edu.get("dates", ""))

    doc.save(docx_path)
    convert(docx_path, pdf_path)

    return pdf_path
